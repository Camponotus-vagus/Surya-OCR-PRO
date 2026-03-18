"""DOCX output writer."""

from __future__ import annotations

import logging
from pathlib import Path

from ..engine.ocr_engine import PageResult
from .writer_base import get_texts_from_results

log = logging.getLogger(__name__)


def write_docx(results: list[PageResult], output_dir: str, pdf_stem: str) -> str:
    """Write all pages to a DOCX file with page breaks.

    Returns the path of the written file.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_BREAK
    except ImportError:
        log.error("python-docx not installed. Run: pip install python-docx")
        raise

    texts = get_texts_from_results(results)
    doc = Document()

    for i, text in enumerate(texts):
        if i > 0:
            # Add page break between pages
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Add page header
        header = doc.add_paragraph()
        run = header.add_run(f"Page {i + 1}")
        run.bold = True
        run.font.size = Pt(14)

        # Add content
        for paragraph_text in text.split("\n\n"):
            paragraph_text = paragraph_text.strip()
            if paragraph_text:
                doc.add_paragraph(paragraph_text)

    filepath = Path(output_dir) / f"{pdf_stem}.docx"
    filepath.parent.mkdir(parents=True, exist_ok=True)

    # Handle file locks
    actual_path = filepath
    for attempt in range(100):
        try:
            doc.save(str(actual_path))
            log.info(f"DOCX written: {actual_path}")
            return str(actual_path)
        except PermissionError:
            actual_path = filepath.parent / f"{pdf_stem}_{attempt + 1}.docx"

    raise PermissionError(f"Cannot write DOCX to {filepath} or alternatives")
